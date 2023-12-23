import os
from docx import Document
from docx.shared import Inches
from datetime import datetime

def process_image(image_path):
    # Process the image and return a suitable representation for docx
    return InlineImage(image_path, width=Inches(4))  # Adjust the width as needed

def combine_markdown_to_docx(output_file, *markdown_files):
    doc = Document()

    for markdown_file in markdown_files:
        markdown_dir = os.path.dirname(markdown_file)
        
        with open(markdown_file, 'r', encoding='utf-8') as f:
            content_lines = f.readlines()

            for line in content_lines:
                if line.strip().startswith("!["):
                    image_path = line.strip().split("](")[1][:-1]
                    full_image_path = os.path.join(markdown_dir, image_path)
                    doc.add_picture(full_image_path, width=Inches(4))  # Adjust the width as needed
                else:
                    doc.add_paragraph(line)

    # Save the DOCX document
    doc.save(output_file)

if __name__ == "__main__":
    markdown_files = [
        '../ED-0001_Example_topic/0001_example.md', 
        #'../ED-0002_Deployment/0001_Deployment.md'
        ]
    current_date = datetime.now().strftime('%Y-%m-%d')
    output_file = f'../System_Docs_{current_date}.docx'

    combine_markdown_to_docx(output_file, *markdown_files)
    print("Markdown files combined into", output_file)

